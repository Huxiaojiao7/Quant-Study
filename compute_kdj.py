#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
绿的谐波 (688017.SH) KDJ 指标计算、可视化与解读报告
"""

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.font_manager import fontManager
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, warnings
warnings.filterwarnings('ignore')

# ============================================================
# 0. 路径配置
# ============================================================
CSV_PATH = r'C:\Users\27924\Desktop\QuantStudy\Task1\绿的谐波_日线数据.csv'
OUTPUT_DIR = r'C:\Users\27924\Desktop\QuantStudy\Task2\outputs'
CHART_PATH = os.path.join(OUTPUT_DIR, 'kdj.png')
DOC_PATH = os.path.join(OUTPUT_DIR, '绿的谐波_688017_KDJ指标解读报告.docx')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ============================================================
# 0. 中文字体
# ============================================================
candidates = ['Microsoft YaHei', 'SimHei', 'KaiTi', 'FangSong', 'STSong']
available = {f.name for f in fontManager.ttflist}
FONT_NAME = None
for c in candidates:
    if c in available:
        FONT_NAME = c
        break
FONT_RC = {'font.family':'sans-serif','font.sans-serif':[FONT_NAME,'DejaVu Sans'],'axes.unicode_minus':False}
plt.rcParams.update(FONT_RC)
print(f"[INFO] 中文字体: {FONT_NAME}")

# ============================================================
# 1. 数据加载
# ============================================================
df = pd.read_csv(CSV_PATH, encoding='utf-8-sig')
df.rename(columns={
    '交易日期':'date','开盘价':'open','最高价':'high','最低价':'low',
    '收盘价':'close','成交量(手)':'volume','成交额(千元)':'amount'
}, inplace=True)
df['date'] = pd.to_datetime(df['date'])
df.sort_values('date', inplace=True)
df.reset_index(drop=True, inplace=True)

ROW_COUNT = len(df)
DATE_MIN = df['date'].min().strftime('%Y-%m-%d')
DATE_MAX = df['date'].max().strftime('%Y-%m-%d')
close_latest = df['close'].iloc[-1]
print(f"[INFO] 数据加载: {ROW_COUNT} 行, {DATE_MIN} ~ {DATE_MAX}")

# ============================================================
# 2. KDJ 计算 (n=9, 经典参数)
# ============================================================
def compute_kdj(df, n=9):
    """计算 KDJ 指标, 返回 K, D, J 序列"""
    low_n = df['low'].rolling(window=n).min()
    high_n = df['high'].rolling(window=n).max()
    rsv = (df['close'] - low_n) / (high_n - low_n) * 100
    rsv = rsv.fillna(50)  # 首段用 50 填充

    k = np.zeros(len(df))
    d = np.zeros(len(df))
    for i in range(len(df)):
        if i == 0:
            k[i] = 50
            d[i] = 50
        else:
            k[i] = 2/3 * k[i-1] + 1/3 * rsv.iloc[i]
            d[i] = 2/3 * d[i-1] + 1/3 * k[i]
    j = 3 * k - 2 * d

    df['KDJ_K'] = k
    df['KDJ_D'] = d
    df['KDJ_J'] = j
    return df

df = compute_kdj(df, n=9)

# ============================================================
# 3. 可视化 (最近 200 日)
# ============================================================
DISPLAY_N = min(200, ROW_COUNT)
df_disp = df.tail(DISPLAY_N).copy()

RED   = '#DC143C'
GREEN = '#228B22'
BLUE  = '#1f77b4'
ORANGE = '#ff7f0e'
PURPLE = '#9467bd'

print("[INFO] 生成 KDJ 可视化...")
fig, ax = plt.subplots(figsize=(16, 5))

ax.plot(df_disp['date'], df_disp['KDJ_K'], color=BLUE,  linewidth=1.2, label='K')
ax.plot(df_disp['date'], df_disp['KDJ_D'], color=ORANGE, linewidth=1.2, label='D')
ax.plot(df_disp['date'], df_disp['KDJ_J'], color=PURPLE, linewidth=0.8, label='J')

# 超买超卖区域
ax.axhline(y=80, color=RED,   linestyle='--', alpha=0.5, linewidth=0.8, label='超买线 (80)')
ax.axhline(y=20, color=GREEN, linestyle='--', alpha=0.5, linewidth=0.8, label='超卖线 (20)')
ax.axhline(y=50, color='gray', linestyle=':',  alpha=0.4, linewidth=0.6)

# 填充
ax.fill_between(df_disp['date'], 80, 100, alpha=0.08, color=RED)
ax.fill_between(df_disp['date'], 0, 20,   alpha=0.08, color=GREEN)

ax.set_ylim(0, 100)
ax.set_ylabel('KDJ')
ax.set_title('绿的谐波 (688017.SH) — KDJ (9, 3, 3)', fontsize=13, fontweight='bold')
ax.legend(loc='upper left', fontsize=9, ncol=6)
ax.grid(True, alpha=0.3)
fig.tight_layout()
fig.savefig(CHART_PATH, dpi=150, bbox_inches='tight')
plt.close(fig)
print(f"[INFO] KDJ 图表已保存: {CHART_PATH}")

# ============================================================
# 4. KDJ 指标解读
# ============================================================
latest = df.iloc[-1]

# 三线当前值
k_now = latest['KDJ_K']
d_now = latest['KDJ_D']
j_now = latest['KDJ_J']

# 角色判断
if k_now > d_now:
    kd_relation = 'K 线在 D 线上方，多头排列'
    kd_signal = '偏多'
else:
    kd_relation = 'K 线在 D 线下方，空头排列'
    kd_signal = '偏空'

# J 线位置
if j_now > 100:
    j_status = f'J 线 = {j_now:.1f}（>100），处于**严重超买**状态，短期回调风险增大'
elif j_now > 80:
    j_status = f'J 线 = {j_now:.1f}（>80），处于**超买**区域，追高需谨慎'
elif j_now < 0:
    j_status = f'J 线 = {j_now:.1f}（<0），处于**严重超卖**状态，技术性反弹概率较高'
elif j_now < 20:
    j_status = f'J 线 = {j_now:.1f}（<20），处于**超卖**区域，关注反弹信号'
elif j_now > 50:
    j_status = f'J 线 = {j_now:.1f}，位于 50 上方偏多区域'
else:
    j_status = f'J 线 = {j_now:.1f}，位于 50 下方偏空区域'

# 金叉死叉检测
recent = df.tail(DISPLAY_N)
k_series = df['KDJ_K']
d_series = df['KDJ_D']

kdj_golden = (k_series > d_series) & (k_series.shift(1) <= d_series.shift(1))
kdj_death  = (k_series < d_series) & (k_series.shift(1) >= d_series.shift(1))

total_golden = kdj_golden.sum()
total_death  = kdj_death.sum()
recent_golden_20 = kdj_golden.tail(20).sum()
recent_death_20  = kdj_death.tail(20).sum()

# J 线极端值
j_above_100 = (df['KDJ_J'] > 100).sum()
j_below_0   = (df['KDJ_J'] < 0).sum()
j_recent_above_100 = ((recent['KDJ_J'] > 100).sum())
j_recent_below_0   = ((recent['KDJ_J'] < 0).sum())

# 钝化检测
recent_j_max = recent['KDJ_J'].max()
recent_j_min = recent['KDJ_J'].min()

if recent_j_max > 95 and (recent['KDJ_J'] > 80).sum() > DISPLAY_N * 0.4:
    passivation = '近 20 日内 J 线持续高位（>80），存在**高位钝化**特征，需警惕趋势可能转折。'
elif recent_j_min < 5 and (recent['KDJ_J'] < 20).sum() > DISPLAY_N * 0.4:
    passivation = '近 20 日内 J 线持续低位（<20），存在**低位钝化**特征，可能酝酿底部反弹。'
else:
    passivation = '未发现明显钝化信号，KDJ 摆动特征正常。'

# 背离检测
recent_close = recent['close']
recent_j = recent['KDJ_J']
close_high_idx = recent_close.idxmax()
j_high_idx = recent_j.idxmax()
close_low_idx = recent_close.idxmin()
j_low_idx = recent_j.idxmin()

divergence_text = ''
if close_high_idx != j_high_idx:
    if recent_close.iloc[-1] > recent_close.iloc[0] and recent_j.iloc[-1] < recent_j.iloc[0]:
        divergence_text = '⚠ 价格走高但 J 线走低，近 20 日内存在**顶背离**信号，短期可能面临回调。'
elif close_low_idx != j_low_idx:
    if recent_close.iloc[-1] < recent_close.iloc[0] and recent_j.iloc[-1] > recent_j.iloc[0]:
        divergence_text = '⚠ 价格走低但 J 线走高，近 20 日内存在**底背离**信号，关注反弹机会。'

# 趋势判断
k_trend = '上升' if recent['KDJ_K'].iloc[-1] > recent['KDJ_K'].iloc[0] else '下降'
d_trend = '上升' if recent['KDJ_D'].iloc[-1] > recent['KDJ_D'].iloc[0] else '下降'
j_trend = '上升' if recent['KDJ_J'].iloc[-1] > recent['KDJ_J'].iloc[0] else '下降'

# 综合解读
kdj_interpretation = (
    f"• **三线定位**：K={k_now:.2f}, D={d_now:.2f}, J={j_now:.2f}。{kd_relation}，{kd_signal}信号。\n"
    f"• **J 线研判**：{j_status}。J 线作为最敏感的先行指标，其极端值往往先于价格给出预警。\n"
    f"• **趋势方向**：K 线近 200 日整体呈{k_trend}趋势，D 线呈{d_trend}趋势，"
    f"短线 J 线呈{j_trend}趋势。\n"
    f"• **金叉/死叉**：全期合计 {total_golden} 次金叉、{total_death} 次死叉；"
    f"最近 20 日 {recent_golden_20} 次金叉、{recent_death_20} 次死叉。\n"
    f"• **极端值**：全期 J 线突破 100 共计 {j_above_100} 次，跌破 0 共计 {j_below_0} 次；"
    f"近 20 日分别出现 {j_recent_above_100} 次（>100）和 {j_recent_below_0} 次（<0）。\n"
    f"• **钝化判断**：{passivation}\n"
)

if divergence_text:
    kdj_interpretation += f"• **背离信号**：{divergence_text}\n"

kdj_interpretation += (
    f"• **综合研判**："
)

if j_now > 100 or k_now > 80:
    kdj_interpretation += (
        "KDJ 三线均处于高位，市场处于极端强势区域。但由于指标已进入超买钝化区，"
        "短线追高风险显著增大。建议等待 K 线下穿 D 线（死叉）或 J 线回落到 80 以下作为短期出场参考。"
        "在强趋势行情中 KDJ 可能持续高位钝化，不宜仅据此做空。"
    )
elif j_now < 0 or k_now < 20:
    kdj_interpretation += (
        "KDJ 三线均处于低位，市场极度超卖。技术性反弹概率在增加，"
        "但仍需等待金叉信号（K 上穿 D）或 J 线回升至 0 以上作为入场确认。"
        "在单边下跌市中，KDJ 低位钝化可能持续较长时间。"
    )
elif k_now > d_now and k_now > 50:
    kdj_interpretation += (
        "KDJ 处于多头排列（K > D），且运行在中轴上方，短期趋势偏强。"
        "可顺势操作，但需关注 J 线是否接近 100 超买区域。"
    )
elif k_now < d_now and k_now < 50:
    kdj_interpretation += (
        "KDJ 处于空头排列（K < D），运行在中轴下方，短期偏弱整理。"
        "关注 J 线是否跌至 0 以下形成超卖。"
    )
else:
    kdj_interpretation += (
        "KDJ 处于震荡区间，方向不明确。建议结合趋势类指标（如 MACD）辅助判断。"
    )

print("[INFO] KDJ 解读完成")

# ============================================================
# 5. 生成 Word 文档
# ============================================================
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_font(cell, size=Pt(9), bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    for p in cell.paragraphs:
        p.alignment = align
        for run in p.runs:
            run.font.size = size
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.bold = bold
            if color:
                run.font.color.rgb = color

def add_styled_table(doc, headers, rows, header_color='1F4E79'):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_font(cell, size=Pt(9), bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        set_cell_shading(cell, header_color)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            set_cell_font(cell, size=Pt(9))
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F7FB')
    return table

def add_highlight_text(doc, text, size=Pt(10.5)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = size
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ============================================================
# 报告正文
# ============================================================

title = doc.add_heading('绿的谐波 (688017.SH) KDJ 指标解读报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph(
    f'分析日期：2026-07-04　　'
    f'数据期间：{DATE_MIN} ~ {DATE_MAX} ({ROW_COUNT} 交易日)　　'
    f'最新收盘：{close_latest:.2f} 元　　'
    f'KDJ 参数：(9, 3, 3)'
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x80,0x80,0x80)

doc.add_paragraph('')
doc.add_paragraph(
    '免责声明：以下内容基于公开数据的技术指标计算与分析，不构成任何投资建议。'
    'KDJ 为摆动类指标，在强趋势行情中存在钝化风险，需结合趋势指标综合判断。'
)
doc.add_paragraph('')

# ---- 一、KDJ 指标简介 ----
doc.add_heading('一、KDJ 指标简介', level=1)

add_highlight_text(doc,
    'KDJ（随机指标 / Stochastic Oscillator）由 George Lane 提出，通过比较收盘价在 '
    '一定周期内价格区间的位置来衡量市场动能的强弱。与 RSI 类似，KDJ 也是摆动类指标，'
    '但由于引入 J 线（反映 K 与 D 的乖离），其灵敏度更高，特别适合捕捉短线转折信号。'
)
doc.add_paragraph('')

add_styled_table(doc,
    ['构成', '计算公式', '特征'],
    [
        ['K 线 (快线)', 'K = 2/3×K_pre + 1/3×RSV', '核心线，反应较灵敏'],
        ['D 线 (慢线)', 'D = 2/3×D_pre + 1/3×K', '信号线，K 的平滑版本'],
        ['J 线 (乖离线)', 'J = 3×K − 2×D', '先行指标，波动最大，可突破 0–100'],
        ['RSV', 'RSV = (C-L_9)/(H_9-L_9)×100', '未成熟随机值，9 日周期'],
    ]
)
doc.add_paragraph('')

doc.add_paragraph(
    '经典判断法则：'
)
notes = [
    'K > 80 / D > 70 / J > 100 → 超买区，短期可能回调',
    'K < 20 / D < 30 / J < 0 → 超卖区，短期可能反弹',
    'K 上穿 D → 金叉买入信号；K 下穿 D → 死叉卖出信号',
    'J 线跌破 0（钝化）后再次上穿 → 较强的底部反转信号',
    'J 线突破 100（钝化）后再次下穿 → 较强的顶部反转信号',
    '在强趋势中，KDJ 可能持续高位/低位钝化，此时应结合趋势指标判断',
]
for n in notes:
    p = doc.add_paragraph(n, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph('')

# ---- 二、当前 KDJ 状态 ----
doc.add_heading('二、当前 KDJ 状态', level=1)

# 颜色标记
def kdj_color(val, low=20, high=80):
    if val > high: return '🔴'
    if val < low: return '🟢'
    return '🟡'

kdj_rows = [
    ['K (快线)', f'{k_now:.2f}', f'{kdj_color(k_now)} {"超买" if k_now>80 else "超卖" if k_now<20 else "中性"}'],
    ['D (慢线)', f'{d_now:.2f}', f'{kdj_color(d_now, 30, 70)} {"偏高" if d_now>70 else "偏低" if d_now<30 else "中性"}'],
    ['J (乖离线)', f'{j_now:.2f}',
     f'🔴 严重超买(>100)' if j_now>100 else f'🔴 超买(>80)' if j_now>80
     else f'🟢 严重超卖(<0)' if j_now<0 else f'🟢 超卖(<20)' if j_now<20 else '🟡 中性'],
    ['K vs D', kd_relation, kd_signal],
    ['近20日 K 趋势', k_trend, '—'],
    ['近20日 D 趋势', d_trend, '—'],
    ['近20日 J 趋势', j_trend, '—'],
    ['近期 J 最高', f'{recent_j_max:.2f}', '—'],
    ['近期 J 最低', f'{recent_j_min:.2f}', '—'],
]
add_styled_table(doc, ['项目', '数值', '状态'], kdj_rows)
doc.add_paragraph('')

# ---- 三、综合解读 ----
doc.add_heading('三、综合解读', level=1)

add_highlight_text(doc, kdj_interpretation)
doc.add_paragraph('')

# ---- 四、交叉信号统计 ----
doc.add_heading('四、金叉/死叉信号统计', level=1)

signal_rows = [
    ['全期金叉', str(total_golden), 'K 上穿 D'],
    ['全期死叉', str(total_death), 'K 下穿 D'],
    ['近20日金叉', str(recent_golden_20), f'最近发生于 {df[kdj_golden]["date"].max().strftime("%Y-%m-%d") if total_golden > 0 else "无"}'],
    ['近20日死叉', str(recent_death_20), f'最近发生于 {df[kdj_death]["date"].max().strftime("%Y-%m-%d") if total_death > 0 else "无"}'],
]
add_styled_table(doc, ['信号类型', '次数', '备注'], signal_rows)

# 列出最近几次交叉
doc.add_paragraph('')
doc.add_paragraph('最近交叉信号记录：')

recent_crosses = df[kdj_golden | kdj_death].tail(10)[['date', 'KDJ_K', 'KDJ_D', 'KDJ_J']]
cross_rows = []
for _, row in recent_crosses.iterrows():
    ctype = '金叉' if row['KDJ_K'] > row['KDJ_D'] else '死叉'
    cross_rows.append([
        row['date'].strftime('%Y-%m-%d'),
        ctype,
        f'{row["KDJ_K"]:.2f}',
        f'{row["KDJ_D"]:.2f}',
        f'{row["KDJ_J"]:.2f}',
    ])
if cross_rows:
    add_styled_table(doc, ['日期', '类型', 'K', 'D', 'J'], cross_rows)
doc.add_paragraph('')

# ---- 五、极端值统计 ----
doc.add_heading('五、J 线极端值统计', level=1)

extreme_rows = [
    ['J > 100 总次数', str(j_above_100), f'占比 {j_above_100/ROW_COUNT*100:.1f}%'],
    ['J < 0 总次数', str(j_below_0), f'占比 {j_below_0/ROW_COUNT*100:.1f}%'],
    ['近20日 J > 100', str(j_recent_above_100), '—'],
    ['近20日 J < 0', str(j_recent_below_0), '—'],
]
add_styled_table(doc, ['项目', '数值', '说明'], extreme_rows)
doc.add_paragraph('')

# ---- 六、与其他指标对照 ----
doc.add_heading('六、KDJ 与其他技术指标的综合视角', level=1)

add_highlight_text(doc,
    'KDJ 属于摆动类指标，与 RSI 有相似之处（均判断超买超卖），但灵敏度更高、钝化风险也更大。'
    '建议在实战中与其他指标交叉验证：'
)
doc.add_paragraph('')

cross_ref = [
    ['RSI', '同属超买超卖类', '若 RSI 与 KDJ 同时发出超买/超卖信号，可信度更高'],
    ['MACD', '趋势类', 'KDJ 给出金叉/死叉时，MACD 是否也处在同向趋势中，可过滤假信号'],
    ['布林带', '波动/趋势类', 'KDJ 超卖 + 价格触及布林下轨 → 反弹概率更大的双重确认'],
    ['ATR', '波动类', '高 ATR 环境下 KDJ 钝化概率大，需降低对其超买超卖信号的权重'],
]
add_styled_table(doc, ['指标', '类型', '联合使用方式'], cross_ref)
doc.add_paragraph('')

# ---- 七、参数说明 ----
doc.add_heading('七、参数与免责', level=1)

add_highlight_text(doc,
    f'本报告 KDJ 参数为 (9, 3, 3)，即计算周期 N=9、EMA 权重满足 K 的平滑系数为 1/3。'
    f'数据来源 Tushare ({DATE_MIN} ~ {DATE_MAX})，共 {ROW_COUNT} 个交易日。'
)
doc.add_paragraph('')

add_highlight_text(doc,
    'KDJ 的局限性：在单边强趋势行情中，J 线可能长时间停留在 >100 或 <0 的钝化区域，'
    '此时金叉/死叉信号频繁出现但多为假信号。请务必结合 MACD、均线等趋势指标综合判断，'
    '不可仅凭 KDJ 单一指标做出交易决策。'
)

doc.add_paragraph('')
p = doc.add_paragraph('— 报告完 —')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 保存
# ============================================================
doc.save(DOC_PATH)
print(f"\n[SUCCESS] KDJ 解读报告已生成: {DOC_PATH}")
