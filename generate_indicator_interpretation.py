#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
绿的谐波 (688017.SH) 技术指标解读报告 — 生成 Word 文档
基于 RSI / MACD / 布林带 / ATR 四项指标结果进行专业解读
"""

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import warnings
warnings.filterwarnings('ignore')

# ============================================================
# 0. 配置
# ============================================================
CSV_PATH = r'C:\Users\27924\Desktop\QuantStudy\Task1\绿的谐波_日线数据.csv'
OUTPUT_PATH = r'C:\Users\27924\Desktop\QuantStudy\Task2\outputs\绿的谐波_688017_技术指标解读报告.docx'

# ============================================================
# 1. 数据加载与指标计算
# ============================================================
df = pd.read_csv(CSV_PATH, encoding='utf-8-sig')
df.rename(columns={
    '交易日期': 'date', '开盘价': 'open', '最高价': 'high',
    '最低价': 'low', '收盘价': 'close', '成交量(手)': 'volume',
    '成交额(千元)': 'amount'
}, inplace=True)
df['date'] = pd.to_datetime(df['date'])
df.sort_values('date', inplace=True)
df.reset_index(drop=True, inplace=True)

# --- RSI ---
def compute_rsi(series, period=14):
    delta = series.diff()
    gain = delta.where(delta > 0, 0.0)
    loss = (-delta).where(delta < 0, 0.0)
    avg_gain = gain.ewm(alpha=1/period, adjust=False).mean()
    avg_loss = loss.ewm(alpha=1/period, adjust=False).mean()
    rs = avg_gain / avg_loss.replace(0, np.nan)
    return 100.0 - (100.0 / (1.0 + rs))

df['rsi_6'] = compute_rsi(df['close'], 6)
df['rsi_14'] = compute_rsi(df['close'], 14)
df['rsi_24'] = compute_rsi(df['close'], 24)

# --- MACD ---
def compute_macd(series, fast=12, slow=26, signal=9):
    ema_fast = series.ewm(span=fast, adjust=False).mean()
    ema_slow = series.ewm(span=slow, adjust=False).mean()
    dif = ema_fast - ema_slow
    dea = dif.ewm(span=signal, adjust=False).mean()
    hist = dif - dea
    return dif, dea, hist

df['dif'], df['dea'], df['macd_hist'] = compute_macd(df['close'])

# --- 布林带 ---
def compute_bollinger(series, period=20, k=2):
    mid = series.rolling(window=period).mean()
    std = series.rolling(window=period).std(ddof=0)
    upper = mid + k * std
    lower = mid - k * std
    bandwidth = (upper - lower) / mid * 100
    pct_b = (series - lower) / (upper - lower)
    return upper, mid, lower, bandwidth, pct_b

df['bb_upper'], df['bb_mid'], df['bb_lower'], df['bb_bandwidth'], df['bb_pct_b'] = \
    compute_bollinger(df['close'])

# --- ATR ---
def compute_atr(df, period=14):
    high, low, close = df['high'], df['low'], df['close']
    prev_close = close.shift(1)
    tr1 = high - low
    tr2 = (high - prev_close).abs()
    tr3 = (low - prev_close).abs()
    tr = pd.concat([tr1, tr2, tr3], axis=1).max(axis=1)
    return tr.ewm(alpha=1/period, adjust=False).mean()

df['atr'] = compute_atr(df, 14)

# 金叉死叉
cross = np.where(
    (df['dif'] > df['dea']) & (df['dif'].shift(1) <= df['dea'].shift(1)), 1,
    np.where(
        (df['dif'] < df['dea']) & (df['dif'].shift(1) >= df['dea'].shift(1)), -1, 0
    )
)
golden_cross = (cross == 1)
death_cross = (cross == -1)

latest = df.iloc[-1]
latest_idx = len(df) - 1

# 最近日期
RECENT_N = 20  # 最近 20 个交易日
recent = df.tail(RECENT_N)
prev_recent = df.tail(40).head(20)

print("[INFO] 指标计算完成")
print(f"最新收盘价: {latest['close']:.2f}")
print(f"RSI-14: {latest['rsi_14']:.2f}")
print(f"MACD DIF/DEA/Hist: {latest['dif']:.4f} / {latest['dea']:.4f} / {latest['macd_hist']:.4f}")
print(f"布林带 %B: {latest['bb_pct_b']:.4f}")
print(f"ATR: {latest['atr']:.4f}")

# ============================================================
# 2. 指标解读分析
# ============================================================

# ---- RSI 解读 ----
rsi_14_latest = latest['rsi_14']
rsi_6_latest = latest['rsi_6']
rsi_24_latest = latest['rsi_24']
rsi_14_recent = recent['rsi_14']
rsi_14_prev = prev_recent['rsi_14']

rsi_overbought = (rsi_14_latest > 70)
rsi_oversold = (rsi_14_latest < 30)
rsi_recent_max = rsi_14_recent.max()
rsi_recent_min = rsi_14_recent.min()
rsi_trend = "上升" if rsi_14_recent.iloc[-1] > rsi_14_recent.iloc[0] else "下降"
rsi_recent_avg = rsi_14_recent.mean()

# RSI 背离检测 (最近 20 日)
price_recent = recent['close']
rsi_high_idx = rsi_14_recent.idxmax()
price_high_idx = price_recent.idxmax()
divergence = ""
if rsi_high_idx != price_high_idx:
    if price_recent.iloc[-1] > price_recent.iloc[0] and rsi_14_recent.iloc[-1] < rsi_14_recent.iloc[0]:
        divergence = "可能存在**顶背离**信号（价格走高而 RSI 走低），需警惕短期回调。"
    elif price_recent.iloc[-1] < price_recent.iloc[0] and rsi_14_recent.iloc[-1] > rsi_14_recent.iloc[0]:
        divergence = "可能存在**底背离**信号（价格走低而 RSI 走高），关注反弹机会。"

# RSI 综合解读文本
if rsi_overbought:
    rsi_level = "超买区 (>70)"
    rsi_interpretation = (
        f"当前 RSI-14 值为 **{rsi_14_latest:.2f}**，处于**超买区域**。"
        f"这意味着近期多头力量已经充分释放，短期可能面临获利回吐压力。"
        f"建议关注是否出现拐头向下信号；若 RSI 从超买区回落至 70 以下，往往是短期调整开始的确认。"
    )
elif rsi_oversold:
    rsi_level = "超卖区 (<30)"
    rsi_interpretation = (
        f"当前 RSI-14 值为 **{rsi_14_latest:.2f}**，处于**超卖区域**。"
        f"价格经过连续下跌后，空方力量趋于衰竭，技术性反弹的概率增大。"
        f"但需配合成交量确认，若放量反弹则可信度更高。"
    )
elif rsi_14_latest > 50:
    rsi_level = "偏多区域 (50~70)"
    rsi_interpretation = (
        f"当前 RSI-14 值为 **{rsi_14_latest:.2f}**，位于 50 中轴上方偏多区域。"
        f"市场整体呈现多头格局，但尚未进入超买极端状态，仍有上行空间。"
        f"短线可参考 RSI-6 的波动捕捉日内级别的买卖点。"
    )
else:
    rsi_level = "偏空区域 (30~50)"
    rsi_interpretation = (
        f"当前 RSI-14 值为 **{rsi_14_latest:.2f}**，位于 50 中轴下方偏空区域。"
        f"市场暂时缺乏上行动能，短期以震荡或偏弱整理为主。"
        f"若 RSI 再度下探至 30 附近而价格未创新低，则构成潜在底部信号。"
    )

# ---- MACD 解读 ----
dif_latest = latest['dif']
dea_latest = latest['dea']
hist_latest = latest['macd_hist']
dif_recent = recent['dif']
dea_recent = recent['dea']

macd_position = "多头区域" if dif_latest > 0 else "空头区域"
hist_direction = "增强" if abs(hist_latest) > abs(recent['macd_hist'].iloc[-2]) else "减弱"

# 金叉死叉
golden_count = golden_cross.sum()
death_count = death_cross.sum()
recent_golden = golden_cross[-RECENT_N:].sum()
recent_death = death_cross[-RECENT_N:].sum()

# 判断近期趋势
dif_trend = "上行" if dif_recent.iloc[-1] > dif_recent.iloc[-5] else ("下行" if dif_recent.iloc[-1] < dif_recent.iloc[-5] else "横盘")

macd_interpretation = (
    f"当前 DIF = **{dif_latest:.4f}**，DEA = **{dea_latest:.4f}**，MACD 柱 = **{hist_latest:.4f}**。\n\n"
    f"• **位置判断**：DIF 位于{'零轴上方' if dif_latest > 0 else '零轴下方'}的{macd_position}，"
    f"{'整体偏多头格局。' if dif_latest > 0 else '空头力量占主导。'}\n"
    f"• **趋势方向**：近 20 日 DIF 整体呈{dif_trend}趋势，"
    f"{'多头动能' if hist_latest > 0 else '空头动能'}正在{hist_direction}。\n"
    f"• **信号统计**：全期间共发生 {int(golden_count)} 次金叉，{int(death_count)} 次死叉；"
    f"最近 20 个交易日内出现 {int(recent_golden)} 次金叉、{int(recent_death)} 次死叉。\n"
    f"• **研判**："
)

if dif_latest > dea_latest and hist_latest > 0:
    macd_interpretation += "DIF 在 DEA 上方运行且柱状图为正值（红柱），当前处于多头主导阶段，建议顺势而为、持仓观察。"
elif dif_latest > dea_latest and hist_latest < 0:
    macd_interpretation += "DIF 仍在 DEA 上方但柱状图转负（绿柱），多头动能衰减，需警惕短期回落。若 DIF 下穿 DEA 将确认死叉信号。"
elif dif_latest < dea_latest and hist_latest < 0:
    macd_interpretation += "DIF 在 DEA 下方运行且柱状图为负值（绿柱），空头占据主导。建议等待金叉信号确认后再考虑参与。"
else:
    macd_interpretation += "DIF 在 DEA 下方但柱状图转正（红柱），空头动能减弱，可能是触底反弹前兆。若 DIF 上穿 DEA 将发出金叉买入信号。"

# 背离检测
price_high_20 = price_recent.max()
dif_high_20 = dif_recent.max()
price_peak_idx = price_recent.idxmax()
dif_peak_idx = dif_recent.idxmax()
if abs(price_peak_idx - dif_peak_idx) > 5:
    if price_recent.iloc[-1] > price_recent.iloc[0] and dif_recent.iloc[-1] < dif_recent.iloc[0]:
        macd_interpretation += "\n\n⚠ 近期价格走高而 DIF 走低，形成**顶背离**信号，需关注可能的趋势反转。"

# ---- 布林带解读 ----
bb_upper_latest = latest['bb_upper']
bb_mid_latest = latest['bb_mid']
bb_lower_latest = latest['bb_lower']
bb_pct_b_latest = latest['bb_pct_b']
bb_bw_latest = latest['bb_bandwidth']
bb_bw_recent = recent['bb_bandwidth']

close_latest = latest['close']

# 布林带位置
if bb_pct_b_latest > 1.0:
    bb_position = "突破上轨"
    bb_position_detail = "价格已突破布林带上轨（%B > 1.0），短线处于极强区域。若成交量配合，可能加速上行；但若无量能支撑，存在回踩上轨确认的需求。"
elif bb_pct_b_latest > 0.8:
    bb_position = "贴近上轨"
    bb_position_detail = "价格贴近布林带上轨运行（%B 在 0.8~1.0），表明多头动能充沛。通常处于上升通道的上沿，短线追高需谨慎。"
elif bb_pct_b_latest > 0.5:
    bb_position = "中轨上方偏强"
    bb_position_detail = "价格运行在中轨与上轨之间偏上位置（%B 在 0.5~0.8），属于健康的上升趋势。中轨（MA20）构成有效支撑。"
elif bb_pct_b_latest > 0.2:
    bb_position = "中轨下方偏弱"
    bb_position_detail = "价格运行在中轨与下轨之间偏下位置（%B 在 0.2~0.5），短期走势偏弱。关注中轨是否能站回，若能则有望转强。"
elif bb_pct_b_latest > 0.0:
    bb_position = "贴近下轨"
    bb_position_detail = "价格贴近布林带下轨（%B 在 0~0.2），短线处于弱势区域。在非单边下跌市中，下轨常构成技术性支撑。"
else:
    bb_position = "跌破下轨"
    bb_position_detail = "价格跌破布林带下轨（%B < 0），短期内跌幅较大。可能出现超跌反弹，但需注意是否存在基本面变化导致的价值中枢下移。"

# 带宽分析
bb_bw_mean = bb_bw_recent.mean()
bb_bw_trend = "扩张" if bb_bw_recent.iloc[-1] > bb_bw_recent.iloc[-10] else "收窄"

bollinger_interpretation = (
    f"当前收盘价 **{close_latest:.2f}** 元，布林带上轨 **{bb_upper_latest:.2f}**，"
    f"中轨 **{bb_mid_latest:.2f}**，下轨 **{bb_lower_latest:.2f}**。\n\n"
    f"• **价格位置**：%B = **{bb_pct_b_latest:.4f}**，处于「**{bb_position}**」。{bb_position_detail}\n"
    f"• **带宽状态**：当前带宽 **{bb_bw_latest:.2f}%**，近 20 日处于{bb_bw_trend}趋势。"
    f"{'波动率放大，趋势行情可能加速运行。' if bb_bw_trend == '扩张' else '波动率收敛，布林带在酝酿方向选择，近期可能出现突破行情。'}\n"
    f"• **关键价位**：中轨（MA20）= {bb_mid_latest:.2f} 元是多空分水岭，"
    f"突破上轨 {bb_upper_latest:.2f} 元需成交量配合，跌破下轨 {bb_lower_latest:.2f} 元则意味着趋势转弱。\n"
    f"• **操作参考**：{'价格突破上轨，可顺势持有；若 %B 从 >1 回落至 <1，可考虑部分止盈。' if bb_pct_b_latest > 1.0 else '在当前位置，可参考中轨与上下轨作为支撑压力来判断短期买卖点。'}"
)

# ---- ATR 解读 ----
atr_latest = latest['atr']
atr_recent = recent['atr']
atr_mean = atr_recent.mean()
atr_pct = atr_latest / close_latest * 100
atr_trend = "增大" if atr_recent.iloc[-1] > atr_recent.mean() else "缩小"

# 计算止损/止盈参考
stop_loss_2x = close_latest - 2 * atr_latest
take_profit_2x = close_latest + 2 * atr_latest
stop_loss_1x = close_latest - atr_latest

atr_interpretation = (
    f"当前 ATR(14) = **{atr_latest:.4f}** 元，占收盘价的 **{atr_pct:.2f}%**。\n\n"
    f"• **波动水平**：近 20 日 ATR 均值 {atr_mean:.2f} 元，当前值"
    f"{'高于' if atr_latest > atr_mean else '低于'}均值，波动率处于{'放大' if atr_latest > atr_mean else '收敛'}状态。\n"
    f"• **波动评估**：ATR/Close = {atr_pct:.2f}%，"
    f"{'属于高波动水平（>5%），短线的多空博弈激烈，单日价格可能剧烈摆动。' if atr_pct > 5 else '属于中等波动水平（2%~5%），市场情绪相对稳定，适合适度仓位参与。' if atr_pct > 2 else '属于低波动水平（<2%），股价走势温和，后续可能出现突破方向的选择。'}\n"
    f"• **止损参考**：若持有多头仓位，基于 2×ATR 的止损价约为 **{stop_loss_2x:.2f}** 元"
    f"（较现价低 {close_latest - stop_loss_2x:.2f} 元）；1×ATR 的较紧止损价约为 **{stop_loss_1x:.2f}** 元。\n"
    f"• **仓位建议**：ATR 偏高→降低仓位以减少回撤风险；ATR 偏低→可适当加大仓位。\n"
    f"• **注意**：ATR 仅衡量波动幅度，不指示涨跌方向。需要结合 MACD 或均线等趋势指标综合判断。"
)

# ============================================================
# 3. 综合研判
# ============================================================

# 综合信号
signals = []
if rsi_14_latest > 50:
    signals.append(("RSI", "偏多"))
else:
    signals.append(("RSI", "偏空"))

if dif_latest > dea_latest and hist_latest > 0:
    signals.append(("MACD", "多头信号"))
elif dif_latest < dea_latest and hist_latest < 0:
    signals.append(("MACD", "空头信号"))
else:
    signals.append(("MACD", "震荡/过渡"))

if bb_pct_b_latest > 0.5:
    signals.append(("布林带", "中轨上方偏强"))
else:
    signals.append(("布林带", "中轨下方偏弱"))

if atr_latest > atr_mean:
    signals.append(("ATR", "波动放大"))
else:
    signals.append(("ATR", "波动收敛"))

bull_count = sum(1 for _, s in signals if '多' in s or '强' in s)
bear_count = sum(1 for _, s in signals if '空' in s or '弱' in s)

if bull_count >= 3:
    overall = "多头共振"
    overall_detail = (
        "多个指标同时发出偏多信号，处于多头共振状态。"
        "短期内趋势向上的概率较大，但需要注意 RSI 是否进入超买区域，"
        "以及 MACD 是否出现动能衰减迹象。建议顺势而为，同时设好止损。"
    )
elif bear_count >= 3:
    overall = "空头共振"
    overall_detail = (
        "多个指标同时指向偏空，处于空头共振状态。"
        "市场整体偏向弱势，需注意仓位控制。"
        "关注布林带下轨和 RSI 超卖区域是否成为短期支撑。"
    )
else:
    overall = "信号分歧"
    overall_detail = (
        "部分指标偏多、部分偏空，多空信号不一致，市场处于方向不明确的震荡格局。"
        "此时单一指标的可信度降低，建议以观望为主，等待多个指标形成共振后再做决策。"
        "可缩小仓位、缩短交易周期以应对不确定性。"
    )

print("[INFO] 解读分析完成")

# ============================================================
# 4. 生成 Word 文档
# ============================================================
doc = Document()

# -- 全局样式 --
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_font(cell, size=Pt(9), bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER):
    for p in cell.paragraphs:
        p.alignment = align
        for run in p.runs:
            run.font.size = size
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            run.font.bold = bold
            if color:
                run.font.color.rgb = color

def add_styled_table(doc, headers, rows, col_widths=None, header_color='1F4E79'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_font(cell, size=Pt(9), bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(cell, header_color)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            set_cell_font(cell, size=Pt(9))
            if r % 2 == 1:
                set_cell_shading(cell, 'F2F7FB')
    return table

def add_highlight_box(doc, text, level=2):
    """添加一个带底色高亮的段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

def add_bold_text(doc, label, content):
    p = doc.add_paragraph()
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run2 = p.add_run(content)
    run2.font.size = Pt(10.5)
    run2.font.name = 'Microsoft YaHei'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# ============================================================
# 报告正文
# ============================================================

title = doc.add_heading('绿的谐波 (688017.SH) 技术指标解读报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')
p = doc.add_paragraph(
    f'分析日期：2026-07-04　　数据期间：2025-07-03 ~ 2026-07-03（{len(df)} 个交易日）　　最新收盘：{close_latest:.2f} 元'
)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(9)
p.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_paragraph('')
doc.add_paragraph(
    '免责声明：以下解读基于公开技术指标计算所得，仅为数据分析结果展示，不构成任何投资建议。'
    '技术指标存在固有滞后性，单一指标不可作为交易决策的唯一依据。请结合基本面与其他分析方法综合判断。'
)

doc.add_paragraph('')

# ========== 一、指标快览 ==========
doc.add_heading('一、核心指标一览', level=1)

snapshot_rows = [
    ['收盘价', f'{close_latest:.2f} 元', '—'],
    ['RSI-6', f'{rsi_6_latest:.2f}', '超短周期，灵敏度高'],
    ['RSI-14', f'{rsi_14_latest:.2f}', '中周期核心参考'],
    ['RSI-24', f'{rsi_24_latest:.2f}', '长周期趋势参考'],
    ['MACD DIF', f'{dif_latest:.4f}', f'{macd_position}'],
    ['MACD DEA', f'{dea_latest:.4f}', '信号线'],
    ['MACD Hist', f'{hist_latest:.4f}', f'{"红柱，多头" if hist_latest > 0 else "绿柱，空头"}动能{hist_direction}'],
    ['布林上轨', f'{bb_upper_latest:.2f} 元', '—'],
    ['布林中轨', f'{bb_mid_latest:.2f} 元', 'MA20，多空分水岭'],
    ['布林下轨', f'{bb_lower_latest:.2f} 元', '—'],
    ['%B', f'{bb_pct_b_latest:.4f}', bb_position],
    ['带宽', f'{bb_bw_latest:.2f}%', f'处于{bb_bw_trend}中'],
    ['ATR(14)', f'{atr_latest:.4f} 元', f'占收盘价 {atr_pct:.2f}%'],
]

add_styled_table(doc, ['指标', '最新值', '状态'], snapshot_rows)
doc.add_paragraph('')

# ========== 二、RSI 解读 ==========
doc.add_heading('二、RSI — 相对强弱指标解读', level=1)

doc.add_heading('2.1 当前状态', level=2)
rsi_detail = [
    ['RSI-6 (短周期)', f'{rsi_6_latest:.2f}'],
    ['RSI-14 (中周期)', f'{rsi_14_latest:.2f}'],
    ['RSI-24 (长周期)', f'{rsi_24_latest:.2f}'],
    ['近20日最高', f'{rsi_recent_max:.2f}'],
    ['近20日最低', f'{rsi_recent_min:.2f}'],
    ['近20日均值', f'{rsi_recent_avg:.2f}'],
    ['近20日趋势', rsi_trend],
    ['所处区域', rsi_level],
]
add_styled_table(doc, ['项目', '数值/状态'], rsi_detail)
doc.add_paragraph('')

doc.add_heading('2.2 综合解读', level=2)
add_highlight_box(doc, rsi_interpretation)
if divergence:
    doc.add_paragraph('')
    p = doc.add_paragraph()
    run = p.add_run(divergence)
    run.bold = True
    run.font.color.rgb = RGBColor(0xDC, 0x14, 0x3C)

doc.add_heading('2.3 多周期对比分析', level=2)
# 计算多周期交叉
rsi_cross_6_14_up = (df['rsi_6'] > df['rsi_14']) & (df['rsi_6'].shift(1) <= df['rsi_14'].shift(1))
rsi_cross_6_14_down = (df['rsi_6'] < df['rsi_14']) & (df['rsi_6'].shift(1) >= df['rsi_14'].shift(1))
rsi_6_above_14 = latest['rsi_6'] > latest['rsi_14']
rsi_14_above_24 = latest['rsi_14'] > latest['rsi_24']

add_bold_text(doc, '短周期(6) vs 中周期(14)：',
    f"{'短周期RSI在中周期上方，短期动量较强。' if rsi_6_above_14 else '短周期RSI在中周期下方，短期动量偏弱。'}")
add_bold_text(doc, '中周期(14) vs 长周期(24)：',
    f"{'中期RSI在长期上方，中期趋势偏多，大方向处于上升通道。' if rsi_14_above_24 else '中期RSI在长期下方，整体趋势偏空。'}")
add_bold_text(doc, '交易信号参考：',
    f"RSI-14 从{'上方' if rsi_14_latest > 50 else '下方'}逼近 50 中轴，"
    f"{'站上 50 则为中线加仓信号。' if rsi_14_latest > 50 else '能否突破 50 是关键转强信号。'}")
doc.add_paragraph('')

# ========== 三、MACD 解读 ==========
doc.add_heading('三、MACD — 指数平滑异同移动平均线解读', level=1)

doc.add_heading('3.1 当前状态', level=2)
macd_detail = [
    ['DIF (12-26 快慢EMA差)', f'{dif_latest:.4f}'],
    ['DEA (DIF的9日EMA)', f'{dea_latest:.4f}'],
    ['MACD柱 (DIF-DEA)', f'{hist_latest:.4f}'],
    ['金叉次数 (全期)', str(int(golden_count))],
    ['死叉次数 (全期)', str(int(death_count))],
    ['近20日金叉', str(int(recent_golden))],
    ['近20日死叉', str(int(recent_death))],
    ['柱状图颜色', '红柱（正值）' if hist_latest > 0 else '绿柱（负值）'],
]
add_styled_table(doc, ['项目', '数值/状态'], macd_detail)
doc.add_paragraph('')

doc.add_heading('3.2 综合解读', level=2)
add_highlight_box(doc, macd_interpretation)
doc.add_paragraph('')

# 金叉死叉历史
if golden_count > 0:
    last_golden = df[golden_cross]['date'].max()
    doc.add_paragraph(f'最近一次金叉发生日期：{last_golden.strftime("%Y-%m-%d")}')
if death_count > 0:
    last_death = df[death_cross]['date'].max()
    doc.add_paragraph(f'最近一次死叉发生日期：{last_death.strftime("%Y-%m-%d")}')

doc.add_paragraph('')

# ========== 四、布林带解读 ==========
doc.add_heading('四、布林带 — Bollinger Bands 解读', level=1)

doc.add_heading('4.1 当前状态', level=2)
bb_detail = [
    ['上轨 (Upper)', f'{bb_upper_latest:.2f} 元', 'MID + 2×STD'],
    ['中轨 (Mid / MA20)', f'{bb_mid_latest:.2f} 元', '多空分界'],
    ['下轨 (Lower)', f'{bb_lower_latest:.2f} 元', 'MID − 2×STD'],
    ['收盘价', f'{close_latest:.2f} 元', '—'],
    ['%B (相对位置)', f'{bb_pct_b_latest:.4f}', bb_position],
    ['带宽 (BandWidth)', f'{bb_bw_latest:.2f}%', f'趋势：{bb_bw_trend}'],
    ['带口宽窄 (上下轨间距)', f'{bb_upper_latest - bb_lower_latest:.2f} 元', '—'],
    ['与上轨距离', f'{bb_upper_latest - close_latest:.2f} 元', '—'],
    ['与中轨距离', f'{close_latest - bb_mid_latest:.2f} 元', '—'],
]
add_styled_table(doc, ['项目', '数值', '说明'], bb_detail)
doc.add_paragraph('')

doc.add_heading('4.2 综合解读', level=2)
add_highlight_box(doc, bollinger_interpretation)
doc.add_paragraph('')

# ========== 五、ATR 解读 ==========
doc.add_heading('五、ATR — 平均真实波幅解读', level=1)

doc.add_heading('5.1 当前状态', level=2)
atr_detail = [
    ['ATR(14)', f'{atr_latest:.4f} 元'],
    ['ATR / Close', f'{atr_pct:.2f}%'],
    ['近20日均值', f'{atr_mean:.2f} 元'],
    ['波动趋势', f'{atr_trend}（{"高于" if atr_latest > atr_mean else "低于"}均值）'],
    ['2×ATR 止损价 (多头)', f'{stop_loss_2x:.2f} 元'],
    ['1×ATR 较紧止损', f'{stop_loss_1x:.2f} 元'],
    ['2×ATR 止盈价 (空头)', f'{take_profit_2x:.2f} 元'],
    ['波动级别', '高' if atr_pct > 5 else ('中' if atr_pct > 2 else '低')],
]
add_styled_table(doc, ['项目', '数值'], atr_detail)
doc.add_paragraph('')

doc.add_heading('5.2 综合解读', level=2)
add_highlight_box(doc, atr_interpretation)
doc.add_paragraph('')

# ========== 六、综合研判 ==========
doc.add_heading('六、综合研判', level=1)

doc.add_heading('6.1 指标共振判断', level=2)
signal_rows = [[name, status] for name, status in signals]
add_styled_table(doc, ['技术指标', '信号方向'], signal_rows)
doc.add_paragraph('')

p = doc.add_paragraph()
run = p.add_run(f'综合评级：{overall}')
run.bold = True
run.font.size = Pt(12)

if overall == "多头共振":
    p.runs[0].font.color.rgb = RGBColor(0xDC, 0x14, 0x3C)
elif overall == "空头共振":
    p.runs[0].font.color.rgb = RGBColor(0x22, 0x8B, 0x22)
else:
    p.runs[0].font.color.rgb = RGBColor(0xFF, 0xA5, 0x00)

doc.add_paragraph('')
add_highlight_box(doc, overall_detail)
doc.add_paragraph('')

doc.add_heading('6.2 参数汇总', level=2)
param_rows = [
    ['RSI', '6 / 14 / 24 (EMA平滑)'],
    ['MACD', 'Fast=12, Slow=26, Signal=9; Hist=DIFF−DEA (不×2)'],
    ['布林带', 'Period=20, K=2'],
    ['ATR', 'Period=14 (EMA平滑)'],
]
add_styled_table(doc, ['指标', '参数设置'], param_rows)
doc.add_paragraph('')

doc.add_heading('6.3 重要提示', level=2)
notes = [
    '技术指标均为滞后指标，反映的是历史价格信息，不可用于预测未来价格。',
    'ATR 不指示方向，仅衡量波动幅度，需搭配趋势类指标使用。',
    '在强趋势行情中，RSI 可能长期处于超买（>70）区域，超买不一定意味立即回调。',
    '布林带在横盘市场中参考价值较高；在强单边行情中，价格可能长时间贴边运行。',
    '本报告所有分析基于 2025-07-03 至 2026-07-03 的日线数据，时效性截至数据末交易日。',
]
for n in notes:
    p = doc.add_paragraph(n, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_paragraph('')
p = doc.add_paragraph('— 报告完 —')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 保存
# ============================================================
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"\n[SUCCESS] 指标解读报告已生成: {OUTPUT_PATH}")
